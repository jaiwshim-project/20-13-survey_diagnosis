#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_background(cell, fill):
    """셀 배경색 설정"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_executive_survey():
    """회장단 본진단 설문 docx 생성"""
    doc = Document()

    # 제목
    title = doc.add_paragraph()
    title_run = title.add_run('👔 회장단 본진단 설문')
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(15, 39, 68)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 부제
    subtitle = doc.add_paragraph('협회의 구조·전략·거버넌스에 관한 진단 설문')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()  # 빈 줄

    # 안내문
    guide = doc.add_paragraph('성실한 응답이 정확한 진단의 기초가 됩니다.')
    guide_run = guide.runs[0]
    guide_run.font.size = Pt(11)
    guide_run.italic = True

    doc.add_paragraph()  # 빈 줄

    sections = [
        {
            'title': 'Ⅰ. 미션·공공성과 전략 작동력',
            'desc': '📌 사전진단 이슈: 미션-실행 괴리',
            'questions': [
                '1. 협회의 미션(전문직 자율성과 공공성 동시 작동)이 실제 의사결정의 최상위 기준으로 작동했던 구체적인 사례는 무엇입니까?',
                '2. 미션 추구 중에 외부 정치적 압력, 이해관계 갈등, 단기 이슈 대응이 우선되었던 경험을 설명해 주십시오.',
                '3. 협회가 추구해야 할 가장 핵심적인 3대 공공 가치는 무엇이라고 생각하십니까?',
                '4. 협회 운영에서 가장 잘 작동되고 있는 전략 또는 정책은 무엇이며, 그 이유는 무엇이라고 생각하십니까?',
                '5. (7점 척도) 협회는 단기 이슈 대응보다 장기 미션 실현에 중점을 두고 운영되고 있다.'
            ]
        },
        {
            'title': 'Ⅱ. 의사결정 구조와 속도',
            'desc': '📌 사전진단 이슈: 의사결정 지연',
            'questions': [
                '6. 중요한 정책/사업 결정은 어떤 절차(회의 횟수, 의견 수렴 방식, 승인 단계)를 거쳐 이루어집니까?',
                '7. 의사결정이 지연되는 주요 원인은 무엇이라고 보십니까? (합의 추구, 정보 부족, 책임 미할당, 권한 불명확 등)',
                '8. 반대로, 신속하고 과감한 결정이 잘 이루어졌던 사례는 무엇이며, 그 조건은 무엇이었습니까?',
                '9. 협회의 합의 중심 문화는 어떤 강점을 가지고 있으며, 이를 어떻게 유지하면서 의사결정 속도를 높일 수 있을까요?',
                '10. (7점 척도) 현재 의사결정 구조는 필요한 속도와 품질의 균형을 이루고 있다.'
            ]
        },
        {
            'title': 'Ⅲ. 거버넌스·연속성·조직기억',
            'desc': '📌 사전진단 이슈: 거버넌스 연속성 단절, 조직기억 소실',
            'questions': [
                '11. 집행부(회장) 교체 시 협회가 직면하는 가장 큰 리스크는 무엇입니까?',
                '12. 과거 리더십 교체 후 정책이나 사업이 중단되거나 크게 변경되었던 구체적인 사례를 설명해 주십시오.',
                '13. 반대로, 리더십이 바뀌어도 지속된 좋은 전통, 제도, 관행이 있다면 무엇이고, 그것이 유지될 수 있는 이유는 무엇입니까?',
                '14. 협회의 과거 의사결정, 정책, 사업 기록(아카이빙)이 얼마나 체계적으로 관리되고 있으며, 신임 리더십에 인수인계되고 있습니까?',
                '15. (7점 척도) 리더십이 바뀌어도 협회의 핵심 전략, 정책, 조직 문화는 일관되게 유지된다.'
            ]
        },
        {
            'title': 'Ⅳ. 역할경계·협업·자율성',
            'desc': '📌 사전진단 이슈: 역할경계 모호',
            'questions': [
                '16. 현재 사무처의 자율성 수준은 어느 정도라고 평가하십니까? (과하지도 부족하지도 않은 적절한 수준이라고 생각하십니까?)',
                '17. 사무처의 실무 판단에 회장단이 개입해야 하는 판단 기준은 무엇이어야 한다고 생각하십니까?',
                '18. 회장단과 사무처의 역할이 가장 잘 협력했던 구체적인 사례는 무엇이며, 그때의 협업 방식은 어떠했습니까?',
                '19. 역할경계가 모호해서 갈등이나 비효율이 발생했던 구체적인 사례를 설명해 주십시오.',
                '20. (7점 척도) 회장단과 사무처 간 역할, 권한, 책임의 경계가 명확하고 일관되게 운영된다.'
            ]
        },
        {
            'title': 'Ⅴ. 갈등·세대·신뢰·공정성',
            'desc': '📌 사전진단 이슈: 반복적 갈등, 세대 간 기준 충돌, 신뢰·공정성 체감 부족',
            'questions': [
                '21. 협회 내에서 주기적으로 반복되는 갈등 유형이 있다면 무엇입니까? (세대 간 기준 충돌, 이해관계 갈등, 권한 경계 문제 등)',
                '22. 과거 갈등을 공식적인 회의·논의로 해결한 성공 사례와 개인적 관계로 남아 있는 갈등 사례를 비교해 주십시오.',
                '23. 회장단 내 세대 간 업무 기준, 의사소통 방식, 가치관의 차이는 구체적으로 어떻게 나타나며, 이것이 협회 운영에 미치는 영향은 무엇입니까?',
                '24. 협회의 의사결정 문화에서 전문성, 데이터, 투명한 절차가 충분히 존중되고 있습니까? 아니면 정치적 고려나 인간관계가 더 중요하다고 느껴집니까?',
                '25. (7점 척도) 협회의 의사결정은 공정하고, 회원과 직원들은 협회를 신뢰한다.'
            ]
        },
        {
            'title': 'Ⅵ. 미래비전·개선방안·민감이슈',
            'desc': '📌 종합평가',
            'questions': [
                '26. 협회의 거버넌스, 의사결정 구조, 역할 체계 중에서 반드시 재설계해야 한다고 생각하는 부분 1가지를 선정하고, 그 이유를 설명해 주십시오.',
                '27. 협회가 직면한 가장 시급한 병목(조직의 성장과 효율성을 가장 크게 제약하는 요인)은 무엇이며, 언제까지 해결해야 합니까?',
                '28. 현실적으로 실행 가능하다고 생각하는 구체적인 개선 방안 3가지를 제안해 주십시오. (각 방안의 비용, 기간, 필요 자원 포함)',
                '29. 3년 후 협회가 지향해야 할 이상적인 운영 모습을 그려주십시오. (의사결정 구조, 조직 문화, 회원·직원 만족도, 공공 가치 실현 등)',
                '30. 이번 조직진단 과정에서 반드시 다루어야 할 민감하거나 핵심적인 이슈가 있다면 적어주십시오. (개인정보 보호 필요한 경우 일반적으로 표기 가능)'
            ]
        }
    ]

    for section in sections:
        # 섹션 제목
        heading = doc.add_paragraph()
        heading_run = heading.add_run(section['title'])
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(15, 39, 68)

        # 섹션 설명
        desc = doc.add_paragraph(section['desc'])
        desc_run = desc.runs[0]
        desc_run.font.size = Pt(10)
        desc_run.font.color.rgb = RGBColor(100, 116, 139)

        # 각 질문
        for question in section['questions']:
            q = doc.add_paragraph(question, style='List Number')
            q_run = q.runs[0]
            q_run.font.size = Pt(11)
            if '척도' in question:
                q_run.font.color.rgb = RGBColor(37, 99, 235)

        doc.add_paragraph()  # 섹션 간 빈 줄

    return doc

def create_employee_survey():
    """직원 본진단 설문 docx 생성"""
    doc = Document()

    # 제목
    title = doc.add_paragraph()
    title_run = title.add_run('👩‍💼 직원(사무처) 본진단 설문')
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(15, 39, 68)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 부제
    subtitle = doc.add_paragraph('협회의 실행·병목·조직문화에 관한 진단 설문')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()  # 빈 줄

    # 안내문
    guide = doc.add_paragraph('완전히 익명으로 처리되며 솔직한 응답을 부탁드립니다.')
    guide_run = guide.runs[0]
    guide_run.font.size = Pt(11)
    guide_run.italic = True

    doc.add_paragraph()  # 빈 줄

    sections = [
        {
            'title': 'Ⅰ. 미션·방향성·프로세스 효율성',
            'desc': '📌 사전진단 이슈: 미션-실행 괴리, 비효율 프로세스',
            'questions': [
                '1. 협회가 추구하는 미션(전문직 자율성과 공공성)을 업무 수행 시에 얼마나 체감하고 있습니까? 구체적인 사례를 들어주십시오.',
                '2. 업무 지시나 우선순위가 자주 바뀌었던 경험이 있다면, 그 영향은 무엇이었습니까?',
                '3. 협회에서 "좋은 일"이라는 기준은 무엇이며, 그 기준이 명확하게 전달되고 있습니까?',
                '4. 협회의 장기적인 전략·계획보다 긴급한 이슈 대응이 우선되는 경우는 얼마나 자주 발생하며, 이것이 업무 효율에 미치는 영향은 무엇입니까?',
                '5. (5점 척도) 협회의 방향성과 우선순위는 일관되게 유지되고 있다.'
            ]
        },
        {
            'title': 'Ⅱ. 의사결정 병목·속도',
            'desc': '📌 사전진단 이슈: 의사결정 병목',
            'questions': [
                '6. 최근 가장 답답했던 의사결정 사례는 무엇이며, 그 과정에서 가장 큰 문제점은 무엇이었습니까?',
                '7. 업무 승인·결재에 몇 단계를 거쳐야 한다고 생각하십니까? 그것이 적절하다고 느끼십니까?',
                '8. 책임자가 불명확해서 업무가 정체되거나 중복되었던 경험이 있다면 설명해 주십시오.',
                '9. 반대로, 의사결정이 빠르고 효과적으로 이루어졌던 사례는 무엇이며, 그 성공 요인은 무엇이었습니까?',
                '10. (5점 척도) 현재 의사결정 구조는 필요한 속도와 품질의 균형을 이루고 있다.'
            ]
        },
        {
            'title': 'Ⅲ. 역할경계·자율성·위임',
            'desc': '📌 사전진단 이슈: 역할경계 모호',
            'questions': [
                '11. 회장단(회장/임원)과 사무처 간의 역할과 권한 경계가 명확하다고 느끼십니까? 모호한 부분이 있다면 무엇입니까?',
                '12. 이미 진행 중인 실무 판단이 상위에서 번복되거나 다시 결정되는 경험이 있다면, 그것이 미친 영향은 무엇입니까?',
                '13. 책임이 주어졌을 때 그에 필요한 권한(예산, 의사결정, 자원 배분)을 충분히 가지고 있다고 느끼십니까? 부족한 부분은 무엇입니까?',
                '14. 부서 간 또는 직급 간 역할이 겹치거나 불분명해서 협업이 어려웠던 경험이 있다면, 그것을 개선할 방법은 무엇이라고 생각하십니까?',
                '15. (5점 척도) 직원으로서 할당된 업무에 대한 자율성과 판단권이 충분하다.'
            ]
        },
        {
            'title': 'Ⅳ. 정보·기록·연속성',
            'desc': '📌 사전진단 이슈: 정보 단절, 조직기억 소실, 거버넌스 연속성',
            'questions': [
                '16. 리더십(회장) 교체 시 사무처 직원으로서 경험했던 가장 큰 어려움은 무엇입니까?',
                '17. 리더십 교체 후 진행 중이던 정책이나 사업이 중단되거나 크게 변경된 구체적인 사례를 설명해 주십시오.',
                '18. 신규 리더십 또는 신규 직원이 업무를 이해하기 위한 인수인계가 체계적으로 이루어지고 있습니까? 부족한 부분은 무엇입니까?',
                '19. 과거 의사결정, 사업 진행 경과, 업무 노하우를 기록하고 공유하는 체계가 충분합니까? 기록 부족으로 어려웠던 구체적인 사례는 무엇입니까?',
                '20. (5점 척도) 리더십이 바뀌어도 협회의 중요한 정책, 업무 흐름, 업무 표준은 일관되게 유지된다.'
            ]
        },
        {
            'title': 'Ⅴ. 평가·세대·문화·신뢰·공정성',
            'desc': '📌 사전진단 이슈: 평가기준 불명확, 세대 간 기준 충돌, 신뢰·공정성 체감 부족',
            'questions': [
                '21. 협회의 리더(회장, 임원)와 직원 간, 또는 세대별로 업무 기준, 의사소통 방식, 가치관의 차이가 있다고 느끼십니까? 구체적인 사례를 들어주십시오.',
                '22. 성과 평가나 인사 대우에서 명확하고 객관적인 기준이 적용되고 있다고 느끼십니까? 아니면 사람에 따라 달라진다고 느끼십니까?',
                '23. 조직에서 업무 개선, 정책 변화에 대한 의견이나 제안을 자유롭게 표현할 수 있다고 느끼십니까? 그렇지 않다면 그 이유는 무엇입니까?',
                '24. 협회 조직 내의 신뢰 수준(회장과 직원, 직원 간, 회의에서의 약속 지킴)은 어느 정도라고 느끼십니까?',
                '25. (5점 척도) 협회의 의사결정과 인사 운영은 공정하고 투명하다.'
            ]
        },
        {
            'title': 'Ⅵ. 미래비전·개선방안·민감이슈',
            'desc': '📌 종합평가',
            'questions': [
                '26. 협회의 업무 기록·매뉴얼·시스템이 얼마나 체계적으로 관리되고 있으며, 신규 직원이나 필요할 때 쉽게 접근할 수 있습니까?',
                '27. 신규 직원이 업무에 적응하는 데 걸리는 시간과 학습 곡선은 어느 정도이며, 더 빠르게 적응시킬 수 있는 방법은 무엇이라고 생각하십니까?',
                '28. 협회의 업무 프로세스에서 반복적으로 비효율이 발생하는 영역은 어디이며, 그 원인과 개선 방법은 무엇입니까?',
                '29. 현실적으로 실행 가능하다고 생각하는 협회 조직·문화·프로세스의 개선 방안 3가지를 제안해 주십시오. (각 방안의 우선순위, 기대 효과 포함)',
                '30. 이번 조직진단 과정에서 반드시 다루어야 할 민감하거나 핵심적인 이슈가 있다면 적어주십시오. (개인정보 보호 필요한 경우 일반적으로 표기 가능)'
            ]
        }
    ]

    for section in sections:
        # 섹션 제목
        heading = doc.add_paragraph()
        heading_run = heading.add_run(section['title'])
        heading_run.font.size = Pt(14)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor(15, 39, 68)

        # 섹션 설명
        desc = doc.add_paragraph(section['desc'])
        desc_run = desc.runs[0]
        desc_run.font.size = Pt(10)
        desc_run.font.color.rgb = RGBColor(100, 116, 139)

        # 각 질문
        for question in section['questions']:
            q = doc.add_paragraph(question, style='List Number')
            q_run = q.runs[0]
            q_run.font.size = Pt(11)
            if '척도' in question:
                q_run.font.color.rgb = RGBColor(16, 185, 129)

        doc.add_paragraph()  # 섹션 간 빈 줄

    return doc

if __name__ == '__main__':
    # 출력 폴더 생성
    output_dir = 'survey_docx'
    os.makedirs(output_dir, exist_ok=True)

    # 회장단 본진단 생성
    exec_doc = create_executive_survey()
    exec_path = os.path.join(output_dir, '회장단본진단.docx')
    exec_doc.save(exec_path)
    print(f'✅ 회장단 본진단 생성 완료: {exec_path}')

    # 직원 본진단 생성
    emp_doc = create_employee_survey()
    emp_path = os.path.join(output_dir, '직원본진단.docx')
    emp_doc.save(emp_path)
    print(f'✅ 직원 본진단 생성 완료: {emp_path}')

    print(f'\n📁 모든 파일이 "{output_dir}" 폴더에 저장되었습니다.')
